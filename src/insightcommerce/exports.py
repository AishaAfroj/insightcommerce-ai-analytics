"""PDF, DOCX, CSV, HTML, and PNG export services for dashboard results."""

from __future__ import annotations

from collections.abc import Iterable
from io import BytesIO

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


class ExportDependencyError(RuntimeError):
    """Raised when optional image-rendering dependencies are unavailable."""


def _safe_rows(frame: pd.DataFrame | None, limit: int = 20) -> list[list[str]]:
    if frame is None or frame.empty:
        return []
    display = frame.head(limit).copy()
    return [list(map(str, display.columns))] + [
        [str(value) for value in row] for row in display.itertuples(index=False, name=None)
    ]


def export_summary_pdf(
    title: str,
    metrics: dict[str, str],
    insights: Iterable[str],
    table_frame: pd.DataFrame | None = None,
) -> bytes:
    """Return a polished compact PDF summary as bytes for a download button."""

    buffer = BytesIO()
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "InsightTitle",
            parent=styles["Title"],
            textColor=colors.HexColor("#1D4ED8"),
            alignment=TA_CENTER,
            fontSize=22,
            leading=27,
        )
    )
    document = SimpleDocTemplate(
        buffer, pagesize=letter, rightMargin=0.65 * inch, leftMargin=0.65 * inch
    )
    story: list = [Paragraph(title, styles["InsightTitle"]), Spacer(1, 0.18 * inch)]
    metric_rows = [[key, value] for key, value in metrics.items()]
    metric_table = Table(metric_rows, colWidths=[2.5 * inch, 4.2 * inch])
    metric_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#DBEAFE")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1E3A8A")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend([metric_table, Spacer(1, 0.22 * inch), Paragraph("Key insights", styles["Heading2"])])
    for insight in insights:
        story.extend([Paragraph(f"• {insight}", styles["BodyText"]), Spacer(1, 5)])
    rows = _safe_rows(table_frame, limit=15)
    if rows:
        story.extend([Spacer(1, 0.15 * inch), Paragraph("Result preview", styles["Heading2"])])
        result_table = Table(rows, repeatRows=1)
        result_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                    ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(result_table)
    document.build(story)
    return buffer.getvalue()


def export_summary_docx(
    title: str,
    metrics: dict[str, str],
    insights: Iterable[str],
    table_frame: pd.DataFrame | None = None,
) -> bytes:
    """Return an editable Word summary as bytes for a download button."""

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(22)
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Shading Accent 1"
    for key, value in metrics.items():
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    document.add_heading("Key insights", level=1)
    for insight in insights:
        document.add_paragraph(insight, style="List Bullet")
    rows = _safe_rows(table_frame, limit=20)
    if rows:
        document.add_heading("Result preview", level=1)
        result = document.add_table(rows=1, cols=len(rows[0]))
        result.style = "Light Shading Accent 1"
        for index, value in enumerate(rows[0]):
            result.rows[0].cells[index].text = value
        for row in rows[1:]:
            cells = result.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_frame_csv(frame: pd.DataFrame) -> bytes:
    """Return a UTF-8 CSV without DataFrame index values."""

    return frame.to_csv(index=False).encode("utf-8")


def export_chart_html(figure: go.Figure) -> bytes:
    """Return a portable interactive chart with Plotly loaded from CDN."""

    return figure.to_html(full_html=True, include_plotlyjs="cdn").encode("utf-8")


def export_chart_png(figure: go.Figure, width: int = 1400, height: int = 850) -> bytes:
    """Return a high-resolution PNG when Kaleido and a Chromium runtime are available."""

    try:
        return figure.to_image(format="png", width=width, height=height, scale=1)
    except Exception as exc:
        raise ExportDependencyError(
            "PNG export needs Kaleido plus an available Chromium/Chrome runtime. "
            "Interactive HTML export remains available."
        ) from exc

